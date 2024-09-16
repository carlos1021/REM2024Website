print("Program starting, grabbing imports ...")
import os
import openai

openai.api_key = os.getenv("OPENAI_API_KEY")

print(f"Grabbed OpenAI API Key: {openai.api_key != None}")

from unstructured.partition.pdf import partition_pdf
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO
import pybase64 as base64
# import PIL
import firebase_admin
from firebase_admin import credentials, storage
import fitz
from PIL import Image
import io
import nltk
# import os
# import openai

nltk.download('punkt_tab')
nltk.download('averaged_perceptron_tagger_eng')

print("Imports complete")

pdf_file_path = './research_paper_example.pdf'

print("Grabbing raw elements")

# openai.api_key = os.getenv["OPENAI_API_KEY"]
MODEL = "gpt-4o"
llm = ChatOpenAI(model=MODEL)

raw_elements = partition_pdf(
    filename=pdf_file_path,
    chunking_strategy="by_title",
    infer_table_structure=True,
    max_characters=1000,
    new_after_n_chars=1500,
    combine_text_undre_n_chars=250,
    strategy="hi_res");

def extract_images_from_pdf(pdf_file_path):
    # The output directory if it doesn't exist
    output_dir = './extracted_images'
    os.makedirs(output_dir, exist_ok=True)

    # Open the PDF file
    pdf_document = fitz.open(pdf_file_path)

    image_count = 0
    # Iterate through each page of the PDF
    for page_num in range(len(pdf_document)):
      page = pdf_document.load_page(page_num)

      # Extract images from the page
      image_list = page.get_images(full=True)

      for img_index, img in enumerate(image_list):
          xref = img[0]
          base_image = pdf_document.extract_image(xref)
          image_bytes = base_image["image"]
          image_ext = base_image["ext"]

          # Convert image to PNG format
          image = Image.open(io.BytesIO(image_bytes))
          image = image.convert("RGB")

          # Save the image
          image_file_path = os.path.join(output_dir, f"image_{page_num + 1}_{img_index + 1}.png")
          image.save(image_file_path, "PNG")
          image_count += 1

    pdf_document.close()
    print(f"Extracted {image_count} images.")

extract_images_from_pdf(pdf_file_path)

tables = []
texts = []

for element in raw_elements:
  if "unstructured.documents.elements.Table" in str(type(element)):
    tables.append(str(element))
  elif "unstructured.documents.elements.CompositeElement" in str(type(element)):
    texts.append(str(element))

print("Appended texts and tables to lists")

def make_prompt(element):
  return f""" You are an assistant tasked with summarizing tables and text from retrieval. \
  These summaries will be embedded and used to retrieve the raw text or table elements. \
  Give a concise summary of the table or text that is well optimized for retrieval. Table or text: {element} """

def generate_text_summaries(texts, tables, summarize_texts=False):
    """
    Summarize text elements

    Args:
        texts: List of str
        tables: List of str
        summarize_texts: Bool to summarize texts
    """

    text_summaries, table_summaries = [], []

    if texts:
        if summarize_texts:
            for text in texts:
                prompt = make_prompt(text)
                completion = openai.chat.completions.create(
                    model=MODEL,
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant tasked with summarizing text."},
                        {"role": "user", "content": prompt}
                    ]
                )
                response = completion.choices[0].message.content
                text_summaries.append(response)

    if tables:
        for table in tables:
            prompt = make_prompt(table)
            completion = openai.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant tasked with summarizing tables."},
                    {"role": "user", "content": prompt}
                ]
            )
            response = completion.choices[0].message.content
            table_summaries.append(response)

    return text_summaries, table_summaries

text_summaries, table_summaries = generate_text_summaries(texts, tables, summarize_texts=False)

print("Generated table summaries")

cred = credentials.Certificate("./REM_Service_Account_Key.json")
firebase_admin.initialize_app(cred, {
    'storageBucket': 'rem2024-f429b.appspot.com'
})

print("Validated Firebase credentials")

def encode_image(image_path):
    # """Encodes an image to a base64 string."""
    # with open(image_path, "rb") as image_file:
    #     return base64.b64encode(image_file.read()).decode("utf-8")
    return None

def upload_image_to_firebase(image_path):
    """Uploads an image to Firebase Storage and returns its public URL."""
    bucket = storage.bucket()
    blob = bucket.blob(os.path.basename(image_path))
    blob.upload_from_filename(image_path)

    # Make the blob publicly viewable.
    blob.make_public()

    # Return the public URL
    return blob.public_url

def analyze_image(img_url):
    """Analyze an image via its Firebase Storage URL and return a summary."""
    prompt = """You are an automotive assistant tasked with summarizing images for retrieval. \
These summaries will be embedded and used to retrieve the raw image. \
Describe concisely the characteristics (shape, color), but do not infer what the image means. \
Only describe the characteristics of the image you see."""

    response = openai.chat.completions.create(
        model=MODEL,
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant."
            },
            {
                "role": "user",
                "content": [
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": img_url,
                        }
                    },
                ],
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        max_tokens=300,
        top_p=0.1
    )

    return response.choices[0].message.content

def generate_image_summaries(image_directory):
    """Generate summaries for images in the specified directory."""
    img_base64_list = []  # Store base64 encoded images
    image_summaries = []  # Store image summaries
    img_url_list = []  # Store Firebase URLs for images

    for filename in sorted(os.listdir(image_directory)):
        if filename.endswith(".png"):
            image_path = os.path.join(image_directory, filename)

            # Convert image to base64 and store in list
            # base64_image = encode_image(image_path)
            # img_base64_list.append(base64_image)
            image_base64_list = None

            # Upload image to Firebase and get the public URL
            img_url = upload_image_to_firebase(image_path)
            img_url_list.append(img_url)

            # Analyze the image using the GPT model with the preserved prompt
            response = analyze_image(img_url)
            image_summaries.append(response)

    return image_summaries, img_base64_list, img_url_list

print("Generating image summaries ...")

image_summaries, img_base64_list, img_url_list = generate_image_summaries('./extracted_images/')
print("Image summaries generated: ")
print("")
# print(image_summaries)

print("Loading in Hugging Face Embeddings")

embeddings = HuggingFaceEmbeddings(
    model_name='BAAI/bge-small-en-v1.5',
    encode_kwargs={'normalize_embeddings': True}
)

documents = []

for i, table in enumerate(tables):
    documents.append({"id": f"table_{i}", "text": table})

for i, text in enumerate(texts):
    documents.append({"id": f"text_{i}", "text": text})

for i, summary in enumerate(image_summaries):
    documents.append({"id": f"image_summary_{i}", "text": summary})

# Create Embeddings and Vector Store
doc_texts = [doc["text"] for doc in documents]  # Extract the text for embedding
vectorstore = FAISS.from_texts(doc_texts, embeddings)

print("Vectorstore created")

document_store = {doc["id"]: doc["text"] for doc in documents}

print("Mapped IDs to Documents (Document Store)")

retriever = vectorstore.as_retriever()

template = """
You are an expert assistant specialized in summarizing research papers into policy briefs.
Use the provided context to create one section of the comprehensive policy brief:

<context>
{context}
</context>

Task: {input}
"""

prompt = ChatPromptTemplate.from_template(template)

# Chain to generate responses from the retrieved documents
doc_chain = create_stuff_documents_chain(llm, prompt)

chain = create_retrieval_chain(retriever, doc_chain)
questions = {
    "Group 1": "Given the research findings, craft an executive summary that highlights the most critical insights in 3-4 sentences. Ensure the summary is concise yet rich in detail, avoiding vague statements and emphasizing the key outcomes of the research. Act like you're presenting to a group.",

    "Group 2": "Provide a detailed background on the issue addressed in the research paper. In 3-4 sentences, explain the context, historical perspective, and the specific problem the paper aims to solve. Include any relevant data or statistics that set the stage for understanding the research. Act like you're presenting to a group.",

    "Group 3": "Identify and summarize the core research question and the associated problem discussed in the paper. Write a 3-4 sentence summary that clearly articulates the research objectives, the problem's significance, and any challenges or gaps the research aims to address. Act like you're presenting to a group.",

    "Group 4": "Outline the key statistical findings of the paper. Provide 3-4 bullet points that include specific numbers, percentages, or other quantitative data that support the research's conclusions. Focus on the most impactful statistics that illustrate the research's outcomes. Act like you're presenting to a group.",

    "Group 5": "Summarize the conclusion and policy recommendations from the research paper in 2-4 sentences. Include specific recommendations backed by data or findings from the paper. Ensure the summary is actionable, with clear suggestions for policy implementation."
}

print("Set up document chain to generate responses from the retrieved documents")

context = ""
for group, question in questions.items():
  response = chain.invoke({"context": context, "input": question})
  if response["answer"]:
    context += f"\n\nAnswer:{response['answer']}"
    print(f"Group: {group}\nAnswer: {response['answer']}\n\n")
  else:
    print(f"Group: {group}\nNo Information\n\n")

print("Final Policy Brief:\n", context)

def filter_relevant_images(image_summaries, img_url_list):
    """
    Filters images based on their relevance to the research paper.

    Args:
    - image_summaries: A list of image summaries generated by GPT.
    - img_url_list: A list of URLs corresponding to the images.

    Returns:
    - A list of URLs of the images deemed relevant to the research paper.
    """
    relevant_images = []
    relevant_urls = []

    # llm = ChatOpenAI(model=MODEL)

    for i, summary in enumerate(image_summaries):
        prompt = f"""
        You are an expert assistant analyzing images in a research paper.
        The following image summary is provided:

        {summary}

        Determine if this image is substantially relevant to the research paper's findings, discussions, or conclusions.
        If the image is merely aesthetic or irrelevant, it should not be included.
        Please answer with "Relevant" or "Not Relevant" and provide a one sentence explanation.
        """

        # Invoke the LLM to determine relevance
        response = llm.invoke(prompt)

        response_text = response.content.strip()

        # Check if the response indicates relevance accurately
        if "Relevant" in response_text and "Not Relevant" not in response_text:
            relevant_images.append(summary)
            relevant_urls.append(img_url_list[i])
            print(f"Image {i+1}: Relevant\nExplanation: {response_text}\n")
        else:
            print(f"Image {i+1}: Not Relevant\nExplanation: {response_text}\n")

    return relevant_urls

relevant_image_urls = filter_relevant_images(image_summaries, img_url_list)

print("Relevant Image URLs:", relevant_image_urls)

def create_policy_brief_document(context, relevant_image_urls, output_path):
    """
    Create a Word document with the policy brief text and relevant images.

    Args:
    - context: The final policy brief text.
    - relevant_image_urls: A list of URLs for the relevant images.
    - output_path: The file path where the document will be saved.

    Returns:
    - None
    """
    doc = Document()

    doc.add_heading('Policy Brief', level=1)
    doc.add_paragraph(context)

    # Add each relevant image to the document
    for img_url in relevant_image_urls:
        # Retrieve the image from the URL
        response = requests.get(img_url)
        img_stream = BytesIO(response.content)

        doc.add_paragraph()

        # Add the image to the document
        doc.add_picture(img_stream, width=Inches(5.0))

    doc.save(output_path)

# Example usage
output_path = './Policy_Brief.docx'
create_policy_brief_document(context, relevant_image_urls, output_path)
print(f"Document saved to {output_path}")