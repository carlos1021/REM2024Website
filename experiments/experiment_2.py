"""GRAPH VISUALIZATION"""
import os
from datasets import load_from_disk
from rouge_score import rouge_scorer
from docx import Document
from bert_score import score as bert_score
from transformers import BartTokenizer, BartForConditionalGeneration
import pandas as pd
import matplotlib.pyplot as plt
import torch

# Define paths for data
original_summary_path = './data/arxiv_summarization_document_50'
llm_brief_path = './data/llm_p_b'  # Folder with generated briefs
generated_brief_path = './data/g_p_b'  # Folder for our model’s generated briefs

# Load original dataset
ds_document = load_from_disk(original_summary_path)

# Initialize ROUGE scorer for ROUGE-L only
rouge_scorer_instance = rouge_scorer.RougeScorer(['rougeL'], use_stemmer=True)

# Function to calculate ROUGE-L
def calculate_rouge_l(original_summary, generated_brief):
    rouge_scores = rouge_scorer_instance.score(original_summary, generated_brief)
    return rouge_scores['rougeL'].fmeasure

# Initialize BERTScore
def calculate_bertscore(original_summary, generated_brief):
    P, R, F1 = bert_score([generated_brief], [original_summary], lang="en")
    return F1.mean().item()

# Initialize BARTScore model
bart_model_name = "facebook/bart-large-cnn"
bart_tokenizer = BartTokenizer.from_pretrained(bart_model_name)
bart_model = BartForConditionalGeneration.from_pretrained(bart_model_name)

def calculate_bartscore(original_summary, generated_brief):
    inputs = bart_tokenizer([original_summary], return_tensors="pt", truncation=True, max_length=512)
    with torch.no_grad():
        outputs = bart_model.generate(inputs.input_ids, max_length=512, num_beams=4, length_penalty=2.0, early_stopping=True)
    generated_text = bart_tokenizer.decode(outputs[0], skip_special_tokens=True)
    return bart_model(inputs.input_ids, labels=outputs).loss.item()

# Function to load generated brief from a DOCX file
def load_generated_brief(filename, folder):
    path = os.path.join(folder, filename)
    doc = Document(path)
    return "\n".join([para.text for para in doc.paragraphs])

# Initialize a list to store evaluation results
results = []

# Evaluate for the first five documents in the dataset
for index in range(5):
    # Load the abstract (reference text)
    original_summary = ds_document[index]['abstract']
    
    # Load generated briefs
    our_generated_brief = load_generated_brief(f"g_p_b_{index}.docx", generated_brief_path)
    gemini_brief = load_generated_brief(f"g_p_b_gemini_{index}.docx", os.path.join(llm_brief_path, 'gemini'))
    openai_brief = load_generated_brief(f"g_p_b_openai_{index}.docx", os.path.join(llm_brief_path, 'openai'))
    
    # Calculate scores for each model's output
    briefs = {
        'Our Model': our_generated_brief,
        'Gemini': gemini_brief,
        'OpenAI': openai_brief
    }
    
    for model_name, brief_text in briefs.items():
        rouge_l = calculate_rouge_l(original_summary, brief_text)
        bert_f1 = calculate_bertscore(original_summary, brief_text)
        bartscore = calculate_bartscore(original_summary, brief_text)
        
        # Append the results for each model
        results.append({
            'Model': model_name,
            'ROUGE-L': rouge_l,
            'BERTScore': bert_f1,
            'BARTScore': bartscore
        })

# Convert results to a DataFrame
df_results = pd.DataFrame(results)

# Calculate average scores for each model and evaluation metric
average_scores = df_results.groupby('Model').mean()

# Plotting the average scores as a bar chart
fig, ax = plt.subplots(figsize=(10, 6))
average_scores.plot(kind='bar', ax=ax, color=['skyblue', 'salmon', 'lightgreen'])

# Adding values on top of each bar
for container in ax.containers:
    ax.bar_label(container, fmt='%.3f', label_type='edge')

# Customize chart
plt.xlabel('Generated Policy Brief')
plt.ylabel('Average Score')
plt.title('Average Evaluation Scores for Our Model, Gemini, and OpenAI')
plt.xticks(rotation=45)
plt.legend(title="Evaluation Metric")
plt.tight_layout()
plt.show()




"""DATA COLLECTION"""
# import os
# from datasets import load_from_disk
# from rouge_score import rouge_scorer
# from nltk.tokenize import word_tokenize
# from docx import Document
# from bert_score import score as bert_score
# from transformers import BartTokenizer, BartForConditionalGeneration
# import openai
# import google.generativeai as genai
# import pandas as pd
# import matplotlib.pyplot as plt
# import torch
# import nltk

# # Ensure required packages are downloaded
# nltk.download('punkt')

# # Set up environment variables for API keys
# openai.api_key = os.getenv("OPENAI_API_KEY")
# genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# # Define paths for data
# original_summary_path = './data/arxiv_summarization_document_50'
# generated_brief_path = './data/g_p_b'
# llm_brief_path = './data/llm_p_b'

# # Ensure output folders exist for saving generated content
# os.makedirs(os.path.join(llm_brief_path, 'gemini'), exist_ok=True)
# os.makedirs(os.path.join(llm_brief_path, 'openai'), exist_ok=True)

# # Load original dataset
# ds_document = load_from_disk(original_summary_path)

# # Initialize ROUGE scorer for ROUGE-L only
# rouge_scorer_instance = rouge_scorer.RougeScorer(['rougeL'], use_stemmer=True)

# # Function to calculate ROUGE-L
# def calculate_rouge_l(original_summary, generated_brief):
#     rouge_scores = rouge_scorer_instance.score(original_summary, generated_brief)
#     return rouge_scores['rougeL'].fmeasure

# # Initialize BERTScore
# def calculate_bertscore(original_summary, generated_brief):
#     P, R, F1 = bert_score([generated_brief], [original_summary], lang="en")
#     return F1.mean().item()

# # Initialize BARTScore model
# bart_model_name = "facebook/bart-large-cnn"
# bart_tokenizer = BartTokenizer.from_pretrained(bart_model_name)
# bart_model = BartForConditionalGeneration.from_pretrained(bart_model_name)

# def calculate_bartscore(original_summary, generated_brief):
#     inputs = bart_tokenizer([original_summary], return_tensors="pt", truncation=True, max_length=512)
#     with torch.no_grad():
#         outputs = bart_model.generate(inputs.input_ids, max_length=512, num_beams=4, length_penalty=2.0, early_stopping=True)
#     generated_text = bart_tokenizer.decode(outputs[0], skip_special_tokens=True)
#     return bart_model(inputs.input_ids, labels=outputs).loss.item()

# # Function to load generated brief from a DOCX file
# def load_generated_brief(filename):
#     path = os.path.join(generated_brief_path, filename)
#     doc = Document(path)
#     return "\n".join([para.text for para in doc.paragraphs])

# # Function to trim text to the first 1000 tokens
# def trim_to_1000_tokens(text):
#     tokens = word_tokenize(text)
#     return " ".join(tokens[:1000])

# # Generate text with Google Gemini API
# def generate_gemini_text(context):
#     prompt = f"""
#     You are an expert assistant specialized in summarizing research papers into policy briefs.
#     Use the provided context to create a comprehensive policy brief: {context}
#     """
#     model = genai.GenerativeModel("gemini-1.5-flash")
#     response = model.generate_content(prompt)
#     return response.text if response else "No response generated."

# # Generate text with OpenAI API
# def generate_openai_text(context):
#     prompt = f"""
#     You are an expert assistant specialized in summarizing research papers into policy briefs.
#     Use the provided context to create a comprehensive policy brief: {context}
#     """
#     response = openai.chat.completions.create(
#         model="gpt-4",
#         messages=[
#             {"role": "system", "content": "You are an expert assistant."},
#             {"role": "user", "content": prompt}
#         ],
#         max_tokens=300,
#         temperature=0.7,
#         top_p=0.1,
#         timeout=30
#     )
#     return response.choices[0].message.content

# # Function to save LLM output as a DOCX file
# def save_as_docx(text, folder, filename):
#     path = os.path.join(llm_brief_path, folder, filename)
#     doc = Document()
#     doc.add_paragraph(text)
#     doc.save(path)
#     return text

# # Function to evaluate scores between original summary and generated brief
# def evaluate_policy_brief(original_summary, generated_brief):
#     rouge_l = calculate_rouge_l(original_summary, generated_brief)
#     bert_f1 = calculate_bertscore(original_summary, generated_brief)
#     bartscore = calculate_bartscore(original_summary, generated_brief)
#     return {'ROUGE-L': rouge_l, 'BERTScore': bert_f1, 'BARTScore': bartscore}

# # Initialize a list to store comparison scores
# results = []

# # Run evaluation for the first five documents in the dataset
# for index in range(5):
#     original_summary = ds_document[index]['abstract']
    
#     # Load our model's generated brief
#     our_generated_filename = f"g_p_b_{index}.docx"
#     our_generated_brief = load_generated_brief(our_generated_filename)
    
#     # Generate policy briefs using Gemini and OpenAI LLMs and save them
#     paper_text = ds_document[index]['article']
#     trimmed_text = trim_to_1000_tokens(paper_text)  # Limit input to 1000 tokens
#     gemini_filename = f"g_p_b_gemini_{index}.docx"
#     openai_filename = f"g_p_b_openai_{index}.docx"
    
#     gemini_brief = save_as_docx(generate_gemini_text(trimmed_text), 'gemini', gemini_filename)
#     openai_brief = save_as_docx(generate_openai_text(trimmed_text), 'openai', openai_filename)
    
#     # Evaluate metrics for our model vs Gemini
#     our_vs_gemini = evaluate_policy_brief(our_generated_brief, gemini_brief)
#     our_vs_gemini['Model Comparison'] = 'Our Model vs Gemini'
#     results.append(our_vs_gemini)
    
#     # Evaluate metrics for our model vs OpenAI
#     our_vs_openai = evaluate_policy_brief(our_generated_brief, openai_brief)
#     our_vs_openai['Model Comparison'] = 'Our Model vs OpenAI'
#     results.append(our_vs_openai)

# # Convert results to a DataFrame
# df_results = pd.DataFrame(results)

# # Calculate average scores for each comparison
# average_scores = df_results.groupby('Model Comparison').mean()

# # Plotting the average scores as a bar chart
# average_scores.plot(kind='bar', figsize=(10, 6), color=['skyblue', 'salmon', 'lightgreen'])
# plt.xlabel('Evaluation Metric')
# plt.ylabel('Average Score')
# plt.title('Average Evaluation Scores for Our Model vs Gemini and OpenAI')
# plt.xticks(rotation=45)
# plt.tight_layout()
# plt.show()
