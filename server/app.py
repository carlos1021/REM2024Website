print("Program starting, grabbing imports ...")
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.chains import create_retrieval_chain
from flask import Flask, request, send_file, jsonify
from docx import Document
from docx.shared import Inches
import requests
from io import BytesIO
import os
import openai
import pybase64 as base64
import firebase_admin
from firebase_admin import credentials, storage
import fitz
from PIL import Image
import io
import json
import nltk
import ssl
from flask_cors import CORS
import sys
import re
# from dotenv import load_dotenv

# load_dotenv()
# sys.setrecursionlimit(5000)  # Increase recursion limit to a higher value
# ssl._create_default_https_context = ssl._create_unverified_context

app = Flask(__name__)
CORS(app, origins=[
    "https://rem2024-f429b.firebaseapp.com", 
    "https://rem2024-f429b.web.app", 
    # "http://localhost:5000",
    "https://rem2024website.onrender.com",
    "https://rem-2024.firebaseapp.com",
    "https://rem-2024.web.app"

],
methods=["GET", "POST", "OPTIONS"],
allow_headers=["Content-Type"])

nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')

print("Imports complete")

openai.api_key = os.getenv("OPENAI_API_KEY")
firebase_key_base64 = os.getenv("FIREBASE_SERVICE_KEY")
firebase_key_json = base64.b64decode(firebase_key_base64).decode('utf-8')
firebase_service_account = json.loads(firebase_key_json)
# with open("REM_Service_Account_Key.json") as f:
#     firebase_service_account = json.load(f)

print(f"Grabbed OpenAI API Key: {openai.api_key != None}")
print(f"Grabbed Firebase Service Account Key: {openai.api_key != None}")

MODEL = "gpt-4o"
llm = ChatOpenAI(model=MODEL)

# Firebase initialization
cred = credentials.Certificate(firebase_service_account)
firebase_admin.initialize_app(cred, {
    'storageBucket': 'rem2024-f429b.appspot.com'
})


def extract_images_and_text_from_pdf(pdf_file_path):
    """
    Extracts text and images from a PDF file using the Fitz (PyMuPDF) library,
    captures multi-line figure/table captions, and removes security-sensitive data like IP addresses.
    """
    pdf_document = fitz.open(pdf_file_path)
    image_urls = []
    text_content = []
    figure_table_captions = []
    bucket = storage.bucket()

    # Regex to detect the start of a figure/table heading
    title_pattern = re.compile(r'^(Fig(?:ure)?\.?|Table)\s*\d+(\.|:)?', re.IGNORECASE)
    # Regex to detect lines containing IP addresses
    ip_address_pattern = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')

    # Iterate through each page
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)

        # Extract text from this page
        page_text = page.get_text("text")

        # Remove any lines containing IP addresses
        filtered_lines = []
        for line in page_text.split('\n'):
            if not ip_address_pattern.search(line):  # Exclude lines with IP addresses
                filtered_lines.append(line)

        # Join filtered lines back into text
        filtered_text = '\n'.join(filtered_lines)
        text_content.append(filtered_text)

        # Split text by lines for figure/table detection
        lines = filtered_text.split('\n')
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]

            # Convert image to PNG
            image_pil = Image.open(io.BytesIO(image_bytes))
            image_pil = image_pil.convert("RGB")

            image_buffer = io.BytesIO()
            image_pil.save(image_buffer, format="PNG")
            image_buffer.seek(0)

            # Upload to Firebase
            image_filename = f"image_{page_num + 1}_{img_index + 1}.png"
            blob = bucket.blob(image_filename)
            blob.upload_from_file(image_buffer, content_type="image/png")

            blob.make_public()
            img_url = blob.public_url
            image_urls.append(img_url)

        # ------------------------------------------
        # Detect figure/table headings and collect multi-line captions
        # ------------------------------------------
        current_caption = None  # Accumulates lines for the "active" figure/table

        for line in lines:
            line_strip = line.strip()

            # Check if this line starts a new figure/table caption
            if title_pattern.match(line_strip):
                # If we were already building a caption, finalize and store it
                if current_caption:
                    figure_table_captions.append(current_caption.strip())

                # Start a new caption with the heading
                current_caption = line_strip

            else:
                # If we're currently in a caption, keep appending lines until we
                # detect the next figure/table heading or finish the page.
                if current_caption is not None:
                    # Decide on a stopping heuristic. For now:
                    # If line is not blank, we'll assume it's a continuation.
                    if line_strip:
                        current_caption += " " + line_strip
                    # If line is blank, you could choose to stop collecting,
                    # but let's keep it simple and just ignore blank lines.

        # End of the page: if there's a caption being built, store it
        if current_caption:
            figure_table_captions.append(current_caption.strip())
            current_caption = None

    pdf_document.close()

    # Return the filtered text, the image URLs, and the figure/table captions
    return text_content, image_urls, figure_table_captions


def process_text(text_content):
    """
    Simulates separation of tables and normal text. Fitz doesn't distinguish between them,
    so we use some heuristics like text structure to identify potential table-like text.
    """
    tables, texts = [], []

    for text in text_content:
        if is_table_like(text):
            tables.append(text)
        else:
            texts.append(text)

    return texts, tables


def is_table_like(text):
    """
    Heuristic to determine if a block of text is table-like.
    If the text has lots of tabular structure (e.g., many lines and columns of numbers), it's likely a table.
    """
    return "\t" in text or " | " in text or "-----" in text

# Function to generate text and table summaries using OpenAI
def make_prompt(element):
    return f"""You are an assistant tasked with summarizing tables and text from retrieval. \
    These summaries will be embedded and used to retrieve the raw text or table elements. \
    Give a concise summary of the table or text that is well optimized for retrieval. Table or text: {element}"""


def generate_text_summaries(texts, tables, summarize_texts=False):
    text_summaries, table_summaries = [], []

    if texts:
        if summarize_texts:
            for text in texts:
                prompt = make_prompt(text)
                completion = openai.chat.completions.create(
                    model=MODEL,
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant tasked with summarizing text."},
                        {"role": "user", "content": prompt}
                    ]
                )
                response = completion.choices[0].message.content
                text_summaries.append(response)

    if tables:
        for table in tables:
            prompt = make_prompt(table)
            completion = openai.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": "You are a helpful assistant tasked with summarizing tables."},
                    {"role": "user", "content": prompt}
                ]
            )
            response = completion.choices[0].message.content
            table_summaries.append(response)

    return text_summaries, table_summaries


# Function to analyze an image URL with OpenAI
def analyze_image(img_url):
    print(f"Image URL: {img_url}")
    print("")

    prompt = """You are an assistant tasked with summarizing images for retrieval. \
    These summaries will be embedded and used to retrieve the raw image. \
    Describe concisely the characteristics (shape, color), but do not infer what the image means. \
    Only describe the characteristics of the image you see."""
    try:
        response = openai.chat.completions.create(
            model=MODEL,
            messages=[
                {
                    "role": "system",
                    "content": "You are a helpful assistant."
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": img_url,
                            }
                        },
                    ],
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            max_tokens=300,
            top_p=0.1,
            timeout=30  # Add a timeout of 15 seconds
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Error occurred during OpenAI API call: {e}")
        return None


# Function to generate image summaries from Firebase URLs
def generate_image_summaries(image_urls):
    image_summaries = []
    for img_url in image_urls:
        response = analyze_image(img_url)
        image_summaries.append(response)
    return image_summaries


# Function to filter relevant images
def filter_relevant_images(image_summaries, img_url_list):
    relevant_images = []
    relevant_urls = []

    for i, summary in enumerate(image_summaries):
        prompt = f"""
        You are an expert assistant analyzing images in a research paper.
        The following image summary is provided:

        {summary}

        Determine if this image is substantially relevant to the research paper's findings, discussions, or conclusions.
        If the image is merely aesthetic or irrelevant, it should not be included.
        Please answer with "Relevant" or "Not Relevant" and provide a one sentence explanation.
        """

        response = openai.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": "You are an expert assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=100,
            top_p=0.1,
            timeout=30
        )

        response_text = response.choices[0].message.content.strip()

        if "Relevant" in response_text and "Not Relevant" not in response_text:
            relevant_images.append(summary)
            relevant_urls.append(img_url_list[i])
            print(f"Image {i+1}: Relevant\nExplanation: {response_text}\n")
        else:
            print(f"Image {i+1}: Not Relevant\nExplanation: {response_text}\n")

    return relevant_urls

def generate_policy_brief_title(context):
    """
    Generates a short, policy-style title for the brief based on the final summarized context.
    This title should reflect the main idea relevant for policymakers,
    but should not be identical to the research paper's original title.
    """
    try:
        response = openai.chat.completions.create(
            model=MODEL,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are a helpful assistant specializing in creating brief, policy-oriented titles. "
                        "Avoid referencing 'research paper' directly. Provide a concise headline reflecting "
                        "the main policy topic or insight relevant for decision-makers."
                    )
                },
                {
                    "role": "user",
                    "content": (
                        "Given the following policy brief content, propose a concise yet descriptive title "
                        "that captures the main idea relevant for policymakers:\n\n"
                        f"{context}"
                    )
                }
            ],
            max_tokens=50,
            temperature=0.7
        )
        # Extract and return the generated title
        title = response.choices[0].message.content.strip()
        return title
    except Exception as e:
        print(f"Error generating policy brief title: {e}")
        # Fallback if OpenAI call fails
        return "Policy Insights on Key Findings"

def make_policy_brief_prompt(context, task):
    """
    Produces a prompt directing the LLM to write as if addressing policymakers,
    avoiding references to 'the research paper' or 'the article' repeatedly.
    """
    return f"""
    You are an expert assistant specialized in crafting authoritative policy briefs
    from research findings. Present the information in a direct, policy-oriented voice,
    avoiding repeated phrases like 'the research paper' or 'this paper.' Instead, 
    summarize the insights in a neutral, informative manner suitable for policymakers.

    Current policy brief content:
    {context}

    New task: {task}
    """

# Function to create a Word document
def create_policy_brief_document(context, relevant_image_urls, output_path, title):
    """
    Create a Word document with the dynamic title and policy brief text,
    converting Markdown-style bold/italic formatting into Word styles.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    # Convert Markdown to Word paragraphs with bold/italic
    # Simple approach: parse line by line, replace markdown tokens **...** and *...*
    # For more robust handling, consider using python-markdown or a 3rd-party library.
    lines = context.split('\n')
    for line in lines:
        # Handle subheadings based on Markdown heading syntax (###, ####, etc.)
        if line.startswith('###'):
            heading_level = line.count('#')  # Count the number of '#' to determine heading level
            heading_text = line.lstrip('#').strip()  # Remove leading '#' and any extra spaces
            doc.add_heading(heading_text, level=min(heading_level, 4))  # Word supports heading levels 1-4
            continue

        # For normal lines, we'll parse for bold (**...**) and italic (*...*)
        p = doc.add_paragraph()
        
        # Parse each line for Markdown formatting tokens
        while line:
            bold_match = re.search(r'\*\*(.+?)\*\*', line)
            italic_match = re.search(r'\*(.+?)\*', line)

            # If no more Markdown tokens are found, add the remainder as plain text
            if not bold_match and not italic_match:
                p.add_run(line)
                break

            # Find which match occurs first in the line
            matches = []
            if bold_match:
                matches.append(('bold', bold_match.start(), bold_match.end(), bold_match.group(1)))
            if italic_match:
                matches.append(('italic', italic_match.start(), italic_match.end(), italic_match.group(1)))

            matches.sort(key=lambda x: x[1])  # Sort by start index
            next_token = matches[0]  # Earliest match
            style_type, start_idx, end_idx, text_inside = next_token

            # Add text preceding the token as normal text
            p.add_run(line[:start_idx])

            # Add the token with style
            run = p.add_run(text_inside)
            if style_type == 'bold':
                run.bold = True
            elif style_type == 'italic':
                run.italic = True

            # Move past this token
            line = line[end_idx:]
    
    # Insert each relevant image into the document
    for img_url in relevant_image_urls:
        response = requests.get(img_url)
        img_stream = BytesIO(response.content)
        doc.add_paragraph()
        doc.add_picture(img_stream, width=Inches(5.0))

    doc.save(output_path)


@app.route('/upload', methods=['POST'])
def upload_pdf():
    print("Request received")

    # Check if the request has the file
    if 'file' not in request.files:
        print("No file provided in the request")
        return jsonify({'error': 'No file provided'}), 400

    # Get the PDF file from the request
    file = request.files['file']
    pdf_file_path = './temp_pdf.pdf'
    print(f"File received: {file.filename}")

    # Save the file locally
    file.save(pdf_file_path)
    print(f"File saved locally at {pdf_file_path}")

    # --- Extract text, images, and figure/table titles from PDF ---
    text_content, image_urls, figure_table_titles = extract_images_and_text_from_pdf(pdf_file_path)
    
    # figure_table_titles = extract_figure_table_titles(pdf_file_path)  # newly added
    print(f"Extracted text content: {len(text_content)} pages")
    print(f"Extracted image URLs: {image_urls}")
    print(f"Extracted figure/table titles: {figure_table_titles}")

    # --- Separate the text and tables ---
    texts, tables = process_text(text_content)
    print(f"Texts found: {len(texts)}, Tables found: {len(tables)}")

    # --- Generate summaries for text/tables (optional) ---
    text_summaries, table_summaries = generate_text_summaries(texts, tables, summarize_texts=False)
    print(f"Generated {len(text_summaries)} text summaries and {len(table_summaries)} table summaries")

    # --- Generate summaries for the extracted images ---
    image_summaries = generate_image_summaries(image_urls)
    print(f"Generated {len(image_summaries)} image summaries")

    # --- Build vectorstore from text, tables, and image summaries ---
    print("Loading Hugging Face embeddings...")
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en-v1.5", encode_kwargs={'normalize_embeddings': True})

    documents = []
    for i, table in enumerate(tables):
        documents.append({"id": f"table_{i}", "text": table})
    for i, text_ in enumerate(texts):
        documents.append({"id": f"text_{i}", "text": text_})
    for i, summary in enumerate(image_summaries):
        documents.append({"id": f"image_summary_{i}", "text": summary})

    print(f"Documents prepared for vector store: {len(documents)}")
    doc_texts = [doc["text"] for doc in documents]
    vectorstore = FAISS.from_texts(doc_texts, embeddings)
    print("Vectorstore created")

    retriever = vectorstore.as_retriever()

    # --- Create the LLM chain with a new 'policy' prompt approach ---
    doc_chain = create_stuff_documents_chain(
        llm,
        ChatPromptTemplate.from_template(
            "You are an expert assistant specialized in crafting policy briefs. {context}\n\nTask: {input}"
        )
    )
    chain = create_retrieval_chain(retriever, doc_chain)

    # --- Define tasks to build continuous, non-Q&A style content ---
    tasks = [
        "Provide a concise executive summary (3-4 sentences) for policymakers.",
        "Offer succinct background on the issue in a policy-oriented tone (3-4 sentences).",
        "Summarize the main focus or objective of this body of work (avoid 'paper' references).",
        "Highlight key numbers or statistics in bullet points (3-4 points).",
        "Conclude with policy recommendations and final insights."
    ]

    # --- Build the final context for the policy brief, removing 'Answer:' and adjusting text ---
    context = ""
    for task in tasks:
        # Use our specialized make_policy_brief_prompt
        custom_prompt = make_policy_brief_prompt(context, task)
        response = chain.invoke({"context": context, "input": custom_prompt})
        if response["answer"]:
            cleaned_answer = response["answer"].replace("Answer:", "").replace("Answer", "")
            context += f"\n\n{cleaned_answer}"

    print("Final concatenated policy brief content:\n", context)

    # --- Filter relevant images ---
    relevant_image_urls = filter_relevant_images(image_summaries, image_urls)
    print(f"Relevant Image URLs: {relevant_image_urls}")

    # --- Generate a dynamic policy-style title ---
    title = generate_policy_brief_title(context)
    print(f"Dynamic policy brief title generated: {title}")

    # --- Create and save the policy brief document ---
    output_path = './Policy_Brief.docx'
    create_policy_brief_document(context, relevant_image_urls, output_path, title)
    print(f"Document saved to {output_path}")

    # --- Clean up the temporary PDF file ---
    os.remove(pdf_file_path)
    print(f"Temporary PDF file {pdf_file_path} removed")

    # --- Read the .docx file into memory and return JSON with base64 data + figure/table titles ---
    with open(output_path, 'rb') as f:
        doc_bytes = f.read()

    doc_b64 = base64.b64encode(doc_bytes).decode('utf-8')

    return jsonify({
        'doc_b64': doc_b64,
        'file_name': 'policy_brief.docx',  # or you can rename dynamically
        'titles': figure_table_titles
    })



# Run the Flask server
if __name__ == '__main__':
    app.run(host='0.0.0.0', port=8080)

