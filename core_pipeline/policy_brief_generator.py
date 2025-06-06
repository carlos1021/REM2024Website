"""
Policy Brief Generation Engine - Step 5 of RAG Pipeline

This module generates professional policy briefs from research evidence using the RAG pipeline.
It transforms academic research into actionable policy documents with professional formatting,
visual elements, and evidence-based recommendations.

Key Features:
- Template-driven professional formatting
- Section-aware content generation (Executive Summary, Background, etc.)
- Evidence integration with proper citations
- Policy-focused language translation
- Visual data presentation (charts, tables, infographics)
- Export to multiple formats (HTML, PDF, DOCX)
- IP-safe visual formatting (no copied diagrams)
"""

import os
import re
import json
import logging
from typing import Dict, List, Tuple, Optional, Any, Union
from dataclasses import dataclass, asdict
from datetime import datetime
from enum import Enum
import base64

import openai
import firebase_admin
from firebase_admin import credentials, storage, db

# Import plotting and document generation libraries
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend
import matplotlib.pyplot as plt
import seaborn as sns
import pandas as pd
from wordcloud import WordCloud
import plotly.graph_objects as go
import plotly.express as px
from plotly.offline import plot
import io

# Document generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# Load environment variables
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Import from our pipeline components
from query_processor import QueryProcessor, ask_question, ResponseResult

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class BriefSection(Enum):
    """Policy brief sections"""
    EXECUTIVE_SUMMARY = "executive_summary"
    BACKGROUND = "background"
    KEY_FINDINGS = "key_findings"
    POLICY_RECOMMENDATIONS = "policy_recommendations"
    SUPPORTING_EVIDENCE = "supporting_evidence"
    REFERENCES = "references"


class OutputFormat(Enum):
    """Output format options"""
    HTML = "html"
    DOCX = "docx"
    PDF = "pdf"
    JSON = "json"


@dataclass
class BriefConfig:
    """Configuration for policy brief generation"""
    title: str
    target_audience: str = "Policy Makers"
    executive_length: str = "medium"  # short, medium, long
    include_charts: bool = True
    include_data_tables: bool = True
    include_wordcloud: bool = True
    color_scheme: str = "professional"  # professional, vibrant, minimal
    max_recommendations: int = 5
    citation_style: str = "policy"  # academic, policy, simple


@dataclass
class BriefSection:
    """Container for a policy brief section"""
    title: str
    content: str
    data_visualizations: List[str]  # Base64 encoded images
    supporting_data: Dict[str, Any]
    citations: List[str]
    confidence_score: float


@dataclass
class PolicyBrief:
    """Complete policy brief container"""
    title: str
    subtitle: str
    date_generated: str
    executive_summary: BriefSection
    background: BriefSection
    key_findings: BriefSection
    policy_recommendations: BriefSection
    supporting_evidence: BriefSection
    references: BriefSection
    metadata: Dict[str, Any]
    visualizations: Dict[str, str]  # Section -> base64 image
    

class PolicyBriefGenerator:
    """
    Professional policy brief generator with visual formatting and evidence integration.
    """
    
    def __init__(self,
                 firebase_bucket: str = "rem2024-f429b.appspot.com",
                 firebase_database_url: str = "https://rem2024-f429b-default-rtdb.firebaseio.com",
                 enable_firebase: bool = True,
                 enable_openai: bool = True):
        """
        Initialize the policy brief generator.
        
        Args:
            firebase_bucket: Firebase storage bucket name
            firebase_database_url: Firebase Realtime Database URL
            enable_firebase: Whether to use Firebase for data access
            enable_openai: Whether to use OpenAI for content generation
        """
        self.firebase_bucket = firebase_bucket
        self.firebase_database_url = firebase_database_url
        self.enable_firebase = enable_firebase
        self.enable_openai = enable_openai
        
        # Initialize OpenAI
        if self.enable_openai:
            openai.api_key = os.getenv("OPENAI_API_KEY")
            if not openai.api_key:
                logger.warning("OPENAI_API_KEY not found - using template-based generation")
                self.enable_openai = False
        
        # Initialize Firebase
        if self.enable_firebase:
            self._init_firebase()
        
        # Initialize query processor
        self.query_processor = QueryProcessor()
        
        # Set up plotting style
        plt.style.use('seaborn-v0_8')
        sns.set_palette("husl")
        
        # Color schemes
        self.color_schemes = {
            'professional': {
                'primary': '#2E4A6B',
                'secondary': '#5D7FA3',
                'accent': '#8FA4C7',
                'background': '#F8F9FA',
                'text': '#2C3E50'
            },
            'vibrant': {
                'primary': '#E74C3C',
                'secondary': '#3498DB',
                'accent': '#F39C12',
                'background': '#ECF0F1',
                'text': '#2C3E50'
            },
            'minimal': {
                'primary': '#34495E',
                'secondary': '#7F8C8D',
                'accent': '#95A5A6',
                'background': '#FFFFFF',
                'text': '#2C3E50'
            }
        }
        
        logger.info("Policy brief generator initialized")
    
    def cleanup(self):
        """Clean up resources and ensure proper termination"""
        try:
            # Close all matplotlib figures
            plt.close('all')
            
            # Clear matplotlib cache
            matplotlib.pyplot.clf()
            matplotlib.pyplot.cla()
            
            logger.info("Policy brief generator cleanup completed")
        except Exception as e:
            logger.warning(f"Cleanup warning: {e}")
    
    def _init_firebase(self):
        """Initialize Firebase connection"""
        try:
            # Check if Firebase is already initialized
            firebase_admin.get_app()
            logger.info("Using existing Firebase app")
        except ValueError:
            try:
                firebase_key_base64 = os.getenv("FIREBASE_SERVICE_KEY")
                if not firebase_key_base64:
                    logger.warning("FIREBASE_SERVICE_KEY not found")
                    self.enable_firebase = False
                    return
                
                firebase_key_json = base64.b64decode(firebase_key_base64).decode('utf-8')
                firebase_service_account = json.loads(firebase_key_json)
                
                cred = credentials.Certificate(firebase_service_account)
                firebase_admin.initialize_app(cred, {
                    'storageBucket': self.firebase_bucket,
                    'databaseURL': self.firebase_database_url
                })
                logger.info("Firebase initialized successfully")
                
            except Exception as e:
                logger.warning(f"Failed to initialize Firebase: {e}")
                self.enable_firebase = False
                return
        
        try:
            self.bucket = storage.bucket()
            self.database = db.reference()
            logger.info("Firebase services ready")
        except Exception as e:
            logger.warning(f"Failed to get Firebase services: {e}")
            self.enable_firebase = False
    
    def generate_policy_brief(self, 
                            config: BriefConfig,
                            vector_db_id: str = None,
                            research_focus: str = None) -> PolicyBrief:
        """
        Generate a complete policy brief from research evidence.
        
        Args:
            config: Configuration for the brief
            vector_db_id: Vector database to search (uses latest if None)
            research_focus: Optional focus area for research queries
            
        Returns:
            Complete PolicyBrief object
        """
        start_time = datetime.now()
        
        try:
            logger.info(f"Generating policy brief: '{config.title}'")
            
            # Step 1: Gather comprehensive evidence
            evidence_base = self._gather_evidence(config, vector_db_id, research_focus)
            
            # Step 2: Generate each section
            sections = {}
            
            # Executive Summary
            sections['executive_summary'] = self._generate_executive_summary(config, evidence_base)
            
            # Background & Context
            sections['background'] = self._generate_background(config, evidence_base)
            
            # Key Findings
            sections['key_findings'] = self._generate_key_findings(config, evidence_base)
            
            # Policy Recommendations
            sections['policy_recommendations'] = self._generate_recommendations(config, evidence_base)
            
            # Supporting Evidence
            sections['supporting_evidence'] = self._generate_supporting_evidence(config, evidence_base)
            
            # References
            sections['references'] = self._generate_references(evidence_base)
            
            # Step 3: Create visualizations
            visualizations = self._create_visualizations(config, evidence_base, sections)
            
            # Step 4: Assemble complete brief
            processing_time = (datetime.now() - start_time).total_seconds()
            
            brief = PolicyBrief(
                title=config.title,
                subtitle=f"Evidence-Based Policy Brief for {config.target_audience}",
                date_generated=datetime.now().strftime("%B %d, %Y"),
                executive_summary=sections['executive_summary'],
                background=sections['background'],
                key_findings=sections['key_findings'],
                policy_recommendations=sections['policy_recommendations'],
                supporting_evidence=sections['supporting_evidence'],
                references=sections['references'],
                metadata={
                    'generation_time_seconds': processing_time,
                    'evidence_sources': len(evidence_base),
                    'total_sections': len(sections),
                    'config': asdict(config),
                    'vector_database': vector_db_id
                },
                visualizations=visualizations
            )
            
            logger.info(f"Policy brief generated successfully in {processing_time:.2f}s")
            
            # Clean up resources
            self.cleanup()
            
            return brief
            
        except Exception as e:
            logger.error(f"Failed to generate policy brief: {e}")
            # Still clean up on error
            self.cleanup()
            raise Exception(f"Policy brief generation failed: {e}")
    
    def _gather_evidence(self, config: BriefConfig, vector_db_id: str, research_focus: str) -> List[ResponseResult]:
        """Gather comprehensive evidence from multiple research queries"""
        
        # Define strategic queries for comprehensive coverage
        base_queries = [
            "What are the main research findings and conclusions?",
            "What methodology was used in this research?",
            "What are the key statistical results and significance?",
            "What are the practical implications of this research?",
            "What limitations or challenges does the research identify?"
        ]
        
        # Add focused queries if research focus is provided
        if research_focus:
            focused_queries = [
                f"What does the research say about {research_focus}?",
                f"What are the outcomes related to {research_focus}?",
                f"What recommendations are made regarding {research_focus}?"
            ]
            base_queries.extend(focused_queries)
        
        evidence_base = []
        
        for query in base_queries:
            try:
                result = ask_question(
                    query, 
                    vector_db_id=vector_db_id,
                    use_hybrid=True,
                    max_results=5
                )
                if result and result.confidence_score > 0.3:  # Filter low-confidence results
                    evidence_base.append(result)
                    logger.debug(f"Added evidence for: {query[:50]}...")
            except Exception as e:
                logger.warning(f"Failed to gather evidence for query '{query}': {e}")
        
        logger.info(f"Gathered {len(evidence_base)} evidence sets from {len(base_queries)} queries")
        return evidence_base
    
    def _generate_executive_summary(self, config: BriefConfig, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate executive summary section"""
        
        if not self.enable_openai:
            return self._generate_template_executive_summary(config, evidence_base)
        
        try:
            # Combine key findings from evidence
            key_points = []
            for result in evidence_base[:3]:  # Use top 3 results
                if result.response and len(result.response) > 100:
                    key_points.append(result.response[:300] + "...")
            
            context = "\n\n".join(key_points)
            
            length_guidance = {
                'short': "in 2-3 concise sentences",
                'medium': "in 1-2 paragraphs (4-6 sentences)",
                'long': "in 2-3 paragraphs (6-10 sentences)"
            }
            
            prompt = f"""
Create a professional executive summary for a policy brief titled "{config.title}" for {config.target_audience}.

Research Context:
{context}

Requirements:
- Write {length_guidance.get(config.executive_length, 'in 1-2 paragraphs')}
- Focus on key findings, implications, and actionable insights
- Use clear, policy-oriented language
- Highlight the most important policy-relevant conclusions
- Avoid technical jargon

The executive summary should capture the essence of the research in a way that busy policymakers can quickly understand the key points and implications.
"""
            
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a policy brief writer specializing in translating research into actionable insights for policymakers."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.3
            )
            
            content = response.choices[0].message.content.strip()
            
            # Extract citations
            citations = []
            for result in evidence_base[:3]:
                if result.citations:
                    citations.extend(result.citations[:2])
            
            return BriefSection(
                title="Executive Summary",
                content=content,
                data_visualizations=[],
                supporting_data={},
                citations=citations[:5],  # Limit citations
                confidence_score=self._calculate_section_confidence(evidence_base[:3])
            )
            
        except Exception as e:
            logger.error(f"OpenAI executive summary generation failed: {e}")
            return self._generate_template_executive_summary(config, evidence_base)
    
    def _generate_background(self, config: BriefConfig, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate background and context section"""
        
        # Look for methodology and context evidence
        context_evidence = []
        for result in evidence_base:
            if any(word in result.query.lower() for word in ['methodology', 'background', 'context', 'study']):
                context_evidence.append(result)
        
        if not context_evidence:
            context_evidence = evidence_base[:2]  # Use first 2 if no specific context found
        
        if self.enable_openai and context_evidence:
            try:
                context = "\n\n".join([r.response[:400] for r in context_evidence[:2]])
                
                prompt = f"""
Create a background and context section for a policy brief titled "{config.title}".

Research Information:
{context}

Requirements:
- Explain the research context and why this topic is important
- Describe the study methodology briefly
- Provide relevant background information for policy understanding
- Keep it concise but informative (2-3 paragraphs)
- Use policy-appropriate language

Focus on information that helps policymakers understand the foundation and credibility of the research.
"""
                
                response = openai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "You are a policy brief writer explaining research context to policymakers."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=600,
                    temperature=0.3
                )
                
                content = response.choices[0].message.content.strip()
                
            except Exception as e:
                logger.error(f"Background generation failed: {e}")
                content = self._generate_template_background(context_evidence)
        else:
            content = self._generate_template_background(context_evidence)
        
        citations = []
        for result in context_evidence:
            if result.citations:
                citations.extend(result.citations[:2])
        
        return BriefSection(
            title="Background & Context",
            content=content,
            data_visualizations=[],
            supporting_data={},
            citations=citations[:5],
            confidence_score=self._calculate_section_confidence(context_evidence)
        )

    def _generate_key_findings(self, config: BriefConfig, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate key findings section with data highlights"""
        
        # Look for results and findings evidence
        findings_evidence = []
        for result in evidence_base:
            if any(word in result.query.lower() for word in ['findings', 'results', 'statistical', 'outcomes']):
                findings_evidence.append(result)
        
        if not findings_evidence:
            findings_evidence = evidence_base[:3]
        
        if self.enable_openai and findings_evidence:
            try:
                context = "\n\n".join([r.response[:500] for r in findings_evidence[:3]])
                
                prompt = f"""
Create a key findings section for a policy brief titled "{config.title}".

Research Evidence:
{context}

Requirements:
- Extract and present the most important research findings
- Focus on quantitative results and statistical significance where available
- Organize findings in a clear, logical order
- Use bullet points or numbered lists for clarity
- Emphasize policy-relevant outcomes
- Include specific data points when available
- Keep language accessible to policymakers

Present findings in a way that clearly shows what the research discovered and why it matters for policy.
"""
                
                response = openai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "You are a policy analyst presenting research findings to decision-makers."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=700,
                    temperature=0.3
                )
                
                content = response.choices[0].message.content.strip()
                
            except Exception as e:
                logger.error(f"Key findings generation failed: {e}")
                content = self._generate_template_findings(findings_evidence)
        else:
            content = self._generate_template_findings(findings_evidence)
        
        # Extract supporting data for visualizations
        supporting_data = self._extract_data_for_visualization(findings_evidence)
        
        citations = []
        for result in findings_evidence:
            if result.citations:
                citations.extend(result.citations[:2])
        
        return BriefSection(
            title="Key Findings",
            content=content,
            data_visualizations=[],
            supporting_data=supporting_data,
            citations=citations[:6],
            confidence_score=self._calculate_section_confidence(findings_evidence)
        )
    
    def _generate_recommendations(self, config: BriefConfig, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate policy recommendations section"""
        
        # Look for implications and recommendations evidence
        rec_evidence = []
        for result in evidence_base:
            if any(word in result.query.lower() for word in ['implications', 'recommendations', 'practical', 'policy']):
                rec_evidence.append(result)
        
        if not rec_evidence:
            rec_evidence = evidence_base[:3]
        
        if self.enable_openai and rec_evidence:
            try:
                context = "\n\n".join([r.response[:400] for r in rec_evidence[:3]])
                
                prompt = f"""
Create policy recommendations for a brief titled "{config.title}" based on research evidence.

Research Evidence:
{context}

Requirements:
- Generate {config.max_recommendations} specific, actionable policy recommendations
- Base recommendations directly on the research findings
- Make recommendations concrete and implementable
- Prioritize recommendations by importance/feasibility
- Use clear, directive language (e.g., "Implement...", "Establish...", "Strengthen...")
- Consider both short-term and long-term actions
- Focus on what policymakers can actually do

Each recommendation should be clearly stated with brief rationale based on the evidence.
"""
                
                response = openai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "You are a policy advisor creating actionable recommendations for government officials."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=800,
                    temperature=0.3
                )
                
                content = response.choices[0].message.content.strip()
                
            except Exception as e:
                logger.error(f"Recommendations generation failed: {e}")
                content = self._generate_template_recommendations(rec_evidence, config.max_recommendations)
        else:
            content = self._generate_template_recommendations(rec_evidence, config.max_recommendations)
        
        citations = []
        for result in rec_evidence:
            if result.citations:
                citations.extend(result.citations[:2])
        
        return BriefSection(
            title="Policy Recommendations",
            content=content,
            data_visualizations=[],
            supporting_data={},
            citations=citations[:5],
            confidence_score=self._calculate_section_confidence(rec_evidence)
        )
    
    def _generate_supporting_evidence(self, config: BriefConfig, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate supporting evidence section with data tables"""
        
        # Compile evidence overview
        evidence_summary = []
        data_points = []
        
        for i, result in enumerate(evidence_base[:5]):
            evidence_summary.append(f"**Evidence Source {i+1}:** {result.response[:200]}...")
            
            # Extract numerical data
            numbers = re.findall(r'\b\d+\.?\d*%?\b', result.response)
            if numbers:
                data_points.extend(numbers[:3])  # Take first 3 numbers
        
        content = f"""
**Research Evidence Overview**

This policy brief is based on comprehensive analysis of {len(evidence_base)} evidence sources from academic research. The following provides an overview of the supporting evidence:

""" + "\n\n".join(evidence_summary)
        
        if data_points:
            content += f"\n\n**Key Data Points Identified:** {', '.join(set(data_points[:10]))}"
        
        # Supporting data for tables
        supporting_data = {
            'evidence_count': len(evidence_base),
            'average_confidence': sum(r.confidence_score for r in evidence_base) / len(evidence_base),
            'data_points': data_points,
            'query_types': [r.response_type for r in evidence_base]
        }
        
        all_citations = []
        for result in evidence_base:
            if result.citations:
                all_citations.extend(result.citations)
        
        return BriefSection(
            title="Supporting Evidence",
            content=content,
            data_visualizations=[],
            supporting_data=supporting_data,
            citations=list(set(all_citations))[:10],  # Unique citations, limit to 10
            confidence_score=supporting_data['average_confidence']
        )
    
    def _generate_references(self, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate references and citations section"""
        
        all_citations = []
        for result in evidence_base:
            if result.citations:
                all_citations.extend(result.citations)
        
        # Remove duplicates and format
        unique_citations = list(set(all_citations))
        
        content = "**Research Sources and Citations**\n\n"
        
        if unique_citations:
            for i, citation in enumerate(unique_citations[:15], 1):  # Limit to 15 citations
                content += f"{i}. {citation}\n"
        else:
            content += "Citations are embedded throughout the brief and reference the original research documents analyzed."
        
        content += f"\n\n**Evidence Quality Assessment**\n"
        content += f"- Total evidence sources analyzed: {len(evidence_base)}\n"
        content += f"- Average confidence score: {sum(r.confidence_score for r in evidence_base) / len(evidence_base):.2f}\n"
        content += f"- Generation date: {datetime.now().strftime('%B %d, %Y')}\n"
        
        return BriefSection(
            title="References & Citations",
            content=content,
            data_visualizations=[],
            supporting_data={},
            citations=unique_citations[:15],
            confidence_score=1.0
        )
    
    def _create_visualizations(self, config: BriefConfig, evidence_base: List[ResponseResult], sections: Dict) -> Dict[str, str]:
        """Create data visualizations for the policy brief"""
        visualizations = {}
        
        try:
            # 1. Evidence confidence chart
            if config.include_charts:
                confidence_chart = self._create_confidence_chart(evidence_base, config.color_scheme)
                if confidence_chart:
                    visualizations['confidence'] = confidence_chart
            
            # 2. Word cloud of key findings
            if config.include_wordcloud:
                wordcloud = self._create_wordcloud(sections, config.color_scheme)
                if wordcloud:
                    visualizations['wordcloud'] = wordcloud
            
            # 3. Data distribution chart
            if config.include_data_tables:
                data_chart = self._create_data_distribution_chart(evidence_base, config.color_scheme)
                if data_chart:
                    visualizations['data_distribution'] = data_chart
            
            logger.info(f"Created {len(visualizations)} visualizations")
            
        except Exception as e:
            logger.warning(f"Visualization creation failed: {e}")
        
        return visualizations
    
    def _create_confidence_chart(self, evidence_base: List[ResponseResult], color_scheme: str) -> str:
        """Create confidence score visualization"""
        try:
            colors = self.color_schemes[color_scheme]
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            # Prepare data
            queries = [f"Query {i+1}" for i in range(len(evidence_base))]
            confidences = [r.confidence_score for r in evidence_base]
            
            # Create bar chart
            bars = ax.bar(queries, confidences, color=colors['primary'], alpha=0.7)
            
            # Customize chart
            ax.set_title('Evidence Confidence Scores', fontsize=16, fontweight='bold', color=colors['text'])
            ax.set_ylabel('Confidence Score', fontsize=12, color=colors['text'])
            ax.set_xlabel('Evidence Sources', fontsize=12, color=colors['text'])
            ax.set_ylim(0, 1.0)
            
            # Add value labels on bars
            for bar, conf in zip(bars, confidences):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height + 0.01,
                       f'{conf:.2f}', ha='center', va='bottom', fontsize=10)
            
            # Style
            ax.grid(True, alpha=0.3)
            ax.set_facecolor(colors['background'])
            fig.patch.set_facecolor('white')
            
            plt.tight_layout()
            
            # Convert to base64
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            
            img_str = base64.b64encode(buffer.getvalue()).decode()
            
            # Properly close everything
            plt.close(fig)
            plt.close('all')
            buffer.close()
            
            return img_str
            
        except Exception as e:
            logger.error(f"Confidence chart creation failed: {e}")
            return ""
    
    def _create_wordcloud(self, sections: Dict, color_scheme: str) -> str:
        """Create word cloud from key findings"""
        try:
            colors = self.color_schemes[color_scheme]
            
            # Combine text from key sections
            text_content = ""
            for section_name in ['key_findings', 'policy_recommendations']:
                if section_name in sections:
                    text_content += " " + sections[section_name].content
            
            if len(text_content) < 50:
                return ""
            
            # Create word cloud
            wordcloud = WordCloud(
                width=800, 
                height=400,
                background_color='white',
                colormap='viridis',
                max_words=50,
                relative_scaling=0.5,
                random_state=42
            ).generate(text_content)
            
            # Plot
            fig, ax = plt.subplots(figsize=(12, 6))
            ax.imshow(wordcloud, interpolation='bilinear')
            ax.axis('off')
            ax.set_title('Key Terms and Concepts', fontsize=16, fontweight='bold', pad=20)
            
            fig.patch.set_facecolor('white')
            plt.tight_layout()
            
            # Convert to base64
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            
            img_str = base64.b64encode(buffer.getvalue()).decode()
            
            # Properly close everything
            plt.close(fig)
            plt.close('all')
            buffer.close()
            
            return img_str
            
        except Exception as e:
            logger.error(f"Word cloud creation failed: {e}")
            return ""
    
    def _create_data_distribution_chart(self, evidence_base: List[ResponseResult], color_scheme: str) -> str:
        """Create data distribution visualization"""
        try:
            colors = self.color_schemes[color_scheme]
            
            # Analyze response types
            response_types = {}
            for result in evidence_base:
                rtype = result.response_type
                response_types[rtype] = response_types.get(rtype, 0) + 1
            
            if len(response_types) < 2:
                return ""
            
            fig, ax = plt.subplots(figsize=(10, 6))
            
            # Create pie chart
            labels = list(response_types.keys())
            sizes = list(response_types.values())
            
            # Use colors from scheme
            chart_colors = [colors['primary'], colors['secondary'], colors['accent']] * 3
            
            wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%', 
                                            colors=chart_colors[:len(labels)],
                                            startangle=90)
            
            ax.set_title('Evidence Type Distribution', fontsize=16, fontweight='bold', color=colors['text'])
            
            fig.patch.set_facecolor('white')
            plt.tight_layout()
            
            # Convert to base64
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            
            img_str = base64.b64encode(buffer.getvalue()).decode()
            
            # Properly close everything
            plt.close(fig)
            plt.close('all')
            buffer.close()
            
            return img_str
            
        except Exception as e:
            logger.error(f"Data distribution chart creation failed: {e}")
            return ""
    
    def export_brief(self, brief: PolicyBrief, output_format: OutputFormat, output_path: str) -> str:
        """Export policy brief to specified format"""
        
        try:
            if output_format == OutputFormat.HTML:
                return self._export_html(brief, output_path)
            elif output_format == OutputFormat.DOCX:
                return self._export_docx(brief, output_path)
            elif output_format == OutputFormat.JSON:
                return self._export_json(brief, output_path)
            else:
                raise ValueError(f"Unsupported output format: {output_format}")
                
        except Exception as e:
            logger.error(f"Export failed: {e}")
            raise Exception(f"Failed to export brief: {e}")
    
    def _export_html(self, brief: PolicyBrief, output_path: str) -> str:
        """Export policy brief as HTML"""
        
        html_template = f"""
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{brief.title}</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #333; max-width: 1200px; margin: 0 auto; padding: 20px; }}
        .header {{ text-align: center; border-bottom: 3px solid #2E4A6B; padding-bottom: 20px; margin-bottom: 30px; }}
        .title {{ color: #2E4A6B; font-size: 2.5em; margin-bottom: 10px; }}
        .subtitle {{ color: #5D7FA3; font-size: 1.2em; margin-bottom: 5px; }}
        .date {{ color: #666; font-style: italic; }}
        .section {{ margin-bottom: 40px; }}
        .section-title {{ color: #2E4A6B; font-size: 1.5em; border-bottom: 2px solid #8FA4C7; padding-bottom: 5px; margin-bottom: 15px; }}
        .executive-summary {{ background: #F8F9FA; padding: 20px; border-left: 5px solid #2E4A6B; margin-bottom: 30px; }}
        .recommendation {{ background: #FFF; border: 1px solid #ddd; padding: 15px; margin: 10px 0; border-radius: 5px; }}
        .citation {{ font-size: 0.9em; color: #666; margin: 5px 0; }}
        .visualization {{ text-align: center; margin: 20px 0; }}
        .confidence {{ background: #E8F4FD; padding: 10px; border-radius: 5px; display: inline-block; }}
        .metadata {{ background: #F5F5F5; padding: 15px; border-radius: 5px; font-size: 0.9em; color: #666; }}
    </style>
</head>
<body>
    <div class="header">
        <h1 class="title">{brief.title}</h1>
        <p class="subtitle">{brief.subtitle}</p>
        <p class="date">{brief.date_generated}</p>
    </div>
    
    <div class="section executive-summary">
        <h2 class="section-title">{brief.executive_summary.title}</h2>
        <p>{brief.executive_summary.content.replace(chr(10), '</p><p>')}</p>
        <div class="confidence">Confidence Score: {brief.executive_summary.confidence_score:.2f}</div>
    </div>
    
    <div class="section">
        <h2 class="section-title">{brief.background.title}</h2>
        <p>{brief.background.content.replace(chr(10), '</p><p>')}</p>
    </div>
    
    <div class="section">
        <h2 class="section-title">{brief.key_findings.title}</h2>
        <div>{brief.key_findings.content.replace(chr(10), '<br>')}</div>
    </div>
    
    <div class="section">
        <h2 class="section-title">{brief.policy_recommendations.title}</h2>
        <div>{brief.policy_recommendations.content.replace(chr(10), '<br>')}</div>
    </div>
    
    <div class="section">
        <h2 class="section-title">{brief.supporting_evidence.title}</h2>
        <p>{brief.supporting_evidence.content.replace(chr(10), '</p><p>')}</p>
    </div>
    
    <div class="section">
        <h2 class="section-title">{brief.references.title}</h2>
        <div>{brief.references.content.replace(chr(10), '<br>')}</div>
    </div>
    
    <div class="section metadata">
        <h3>Generation Metadata</h3>
        <p>Generated in {brief.metadata['generation_time_seconds']:.2f} seconds</p>
        <p>Evidence sources: {brief.metadata['evidence_sources']}</p>
        <p>Vector database: {brief.metadata['vector_database']}</p>
    </div>
</body>
</html>
"""
        
        # Add visualizations
        if brief.visualizations:
            viz_html = "<div class='section'><h2 class='section-title'>Data Visualizations</h2>"
            for name, img_data in brief.visualizations.items():
                if img_data:
                    viz_html += f'<div class="visualization"><img src="data:image/png;base64,{img_data}" alt="{name}" style="max-width: 100%; height: auto;"></div>'
            viz_html += "</div>"
            
            # Insert before metadata
            html_template = html_template.replace('<div class="section metadata">', viz_html + '<div class="section metadata">')
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
        
        logger.info(f"HTML export completed: {output_path}")
        return output_path
    
    def _export_docx(self, brief: PolicyBrief, output_path: str) -> str:
        """Export policy brief as Word document"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading(brief.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph(brief.subtitle)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_para = doc.add_paragraph(brief.date_generated)
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # Add sections
        sections = [
            brief.executive_summary,
            brief.background,
            brief.key_findings,
            brief.policy_recommendations,
            brief.supporting_evidence,
            brief.references
        ]
        
        for section in sections:
            doc.add_heading(section.title, level=1)
            doc.add_paragraph(section.content)
            
            if section.citations:
                doc.add_paragraph("References:", style='Heading 3')
                for citation in section.citations[:5]:
                    doc.add_paragraph(citation, style='List Bullet')
            
            doc.add_paragraph()  # Add space
        
        doc.save(output_path)
        logger.info(f"DOCX export completed: {output_path}")
        return output_path
    
    def _export_json(self, brief: PolicyBrief, output_path: str) -> str:
        """Export policy brief as JSON"""
        
        brief_dict = asdict(brief)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(brief_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"JSON export completed: {output_path}")
        return output_path
    
    # Template-based fallback methods
    def _generate_template_executive_summary(self, config: BriefConfig, evidence_base: List[ResponseResult]) -> BriefSection:
        """Generate template-based executive summary when OpenAI is not available"""
        content = f"""
Based on comprehensive analysis of research evidence, this brief presents key findings related to {config.title}. 
The research provides valuable insights for {config.target_audience.lower()} based on {len(evidence_base)} evidence sources. 
Key findings indicate significant implications for policy development and implementation.
"""
        return BriefSection(
            title="Executive Summary",
            content=content,
            data_visualizations=[],
            supporting_data={},
            citations=[],
            confidence_score=0.7
        )
    
    def _generate_template_background(self, evidence: List[ResponseResult]) -> str:
        """Generate template background section"""
        return f"""
This policy brief is based on analysis of research evidence from {len(evidence)} sources. 
The research provides important context for understanding the topic and its policy implications. 
The methodology and approach used in the research ensures reliable and credible findings for policy consideration.
"""
    
    def _generate_template_findings(self, evidence: List[ResponseResult]) -> str:
        """Generate template findings section"""
        content = "**Key Research Findings:**\n\n"
        for i, result in enumerate(evidence[:3], 1):
            content += f"{i}. {result.response[:150]}...\n\n"
        return content
    
    def _generate_template_recommendations(self, evidence: List[ResponseResult], max_recs: int) -> str:
        """Generate template recommendations section"""
        content = "**Policy Recommendations:**\n\n"
        for i in range(min(max_recs, 3)):
            content += f"{i+1}. Implement evidence-based policies addressing the key findings from this research.\n"
        return content
    
    # Utility methods
    def _calculate_section_confidence(self, evidence: List[ResponseResult]) -> float:
        """Calculate confidence score for a section"""
        if not evidence:
            return 0.5
        return sum(r.confidence_score for r in evidence) / len(evidence)
    
    def _extract_data_for_visualization(self, evidence: List[ResponseResult]) -> Dict[str, Any]:
        """Extract numerical data from evidence for visualization"""
        data = {
            'numbers': [],
            'percentages': [],
            'statistics': []
        }
        
        for result in evidence:
            # Extract numbers
            numbers = re.findall(r'\b\d+\.?\d*\b', result.response)
            data['numbers'].extend(numbers[:5])
            
            # Extract percentages
            percentages = re.findall(r'\b\d+\.?\d*%\b', result.response)
            data['percentages'].extend(percentages[:3])
            
            # Extract statistical terms
            stats = re.findall(r'\b(p\s*[<>=]\s*0\.\d+|r\s*=\s*0\.\d+)\b', result.response, re.IGNORECASE)
            data['statistics'].extend(stats[:3])
        
        return data


# Convenience functions
def generate_policy_brief(title: str, 
                         research_focus: str = None,
                         vector_db_id: str = None,
                         target_audience: str = "Policy Makers",
                         output_format: str = "html",
                         output_path: str = None) -> str:
    """
    Convenience function to generate a policy brief.
    
    Args:
        title: Title for the policy brief
        research_focus: Optional focus area for research
        vector_db_id: Vector database ID (uses latest if None)
        target_audience: Target audience for the brief
        output_format: Output format (html, docx, json)
        output_path: Output file path (auto-generated if None)
        
    Returns:
        Path to generated file
    """
    generator = PolicyBriefGenerator()
    
    # Create configuration
    config = BriefConfig(
        title=title,
        target_audience=target_audience,
        executive_length="medium",
        include_charts=True,
        include_data_tables=True,
        include_wordcloud=True,
        color_scheme="professional",
        max_recommendations=5
    )
    
    # Generate brief
    brief = generator.generate_policy_brief(config, vector_db_id, research_focus)
    
    # Determine output path
    if output_path is None:
        safe_title = re.sub(r'[^\w\s-]', '', title).strip()
        safe_title = re.sub(r'[\s_-]+', '_', safe_title)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"policy_brief_{safe_title}_{timestamp}.{output_format}"
    
    # Export
    format_enum = OutputFormat(output_format.lower())
    result = generator.export_brief(brief, format_enum, output_path)
    
    # Clean up and ensure termination
    generator.cleanup()
    
    return result 